"""Generate full Word report. Run: python generate_project_report.py"""
from pathlib import Path
import csv
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = Path(__file__).resolve().parent
OUTPUT = ROOT / "Lattice_EDM_Project_Report.docx"

IMAGES = [
    ("PROBLEM STATEMENT 1 LATTICE STRUICTURE.jpeg", "Figure 1: Lattice structure, 900 um tool, reference (0,0)"),
    ("LATTICE PROBLEM STATEMENT 2.jpeg", "Figure 2: Unknown tool position — Phase 1"),
    ("SHOWINGIMAGES LATTICE CORRECT IMAGES.jpeg", "Figure 3: Favourable vs unfavourable outcomes"),
    ("ACTUAL OUTPUT WE WANT.jpeg", "Figure 4: Target — black supporting ring circular"),
    ("MEASURMEBN.jpeg", "Figure 5: Pore diameter 235.6 um calculation"),
    ("llm project ideation page 2.jpeg", "Figure 6: Project overview — 500 um unit cell, grid logic"),
    ("llm project 3 ideation page 1.jpeg", "Figure 7: Grid subdivision and ML circularity map idea"),
    ("ACTUAL DATSET .png", "Figure 8: Lattice unit cell and 16-run data table"),
    ("ACTUAL IMAGE OF THE 16 DATASETS .png", "Figure 9: SEM images runs 1-16"),
]


def add_image(doc, fname, caption, w=5.5):
    p = ROOT / fname
    if p.exists():
        doc.add_picture(str(p), width=Inches(w))
        c = doc.add_paragraph(caption)
        c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        doc.add_paragraph(f"[Missing: {fname}]")
    doc.add_paragraph()


def add_csv_table(doc, path, title):
    doc.add_paragraph(title).runs[0].bold = True
    with open(path, newline="", encoding="utf-8") as f:
        rows = list(csv.reader(f))
    if not rows:
        return
    t = doc.add_table(rows=len(rows), cols=len(rows[0]))
    t.style = "Table Grid"
    for i, row in enumerate(rows):
        for j, v in enumerate(row):
            t.rows[i].cells[j].text = str(v)
    doc.add_paragraph()


def build():
    doc = Document()
    doc.add_heading("Lattice Structure EDM Project — Full Report", 0)
    doc.add_paragraph(
        "Phase 1 & 2 analysis, geometry (pore 235.6 um, tool 900 um), "
        "16 lab runs, SEM validation, and final EDM recommendations."
    )
    doc.add_page_break()

    doc.add_heading("1. Introduction", 1)
    doc.add_paragraph(
        "Goal: nearly circular machined pore with continuous supporting-material (black) boundary. "
        "Nodes may be destroyed. Tool 900 um; pore 235.6 um; unit cell 500 um; tool/pore ratio 3.82."
    )

    doc.add_heading("2. Geometry", 1)
    add_csv_table(doc, ROOT / "data" / "lattice_geometry.csv", "Lattice geometry")
    add_image(doc, "MEASURMEBN.jpeg", IMAGES[4][1], 4.5)

    doc.add_heading("3. Problem diagrams", 1)
    for fname, cap in IMAGES[:4]:
        add_image(doc, fname, cap, 5.0)

    doc.add_heading("4. LLM / grid ideation", 1)
    doc.add_paragraph(
        "Phase 2 working area: 3x3 unit cells = 1500 x 1500 um (tool exceeds single 500 um cell). "
        "Grid points (a,b) map tool landing positions for circularity prediction."
    )
    add_image(doc, "llm project ideation page 2.jpeg", IMAGES[5][1], 5.0)
    add_image(doc, "llm project 3 ideation page 1.jpeg", IMAGES[6][1], 5.0)

    doc.add_heading("5. Experimental data", 1)
    add_image(doc, "ACTUAL DATSET .png", IMAGES[7][1], 6.0)
    add_csv_table(doc, ROOT / "data" / "original_16_runs.csv", "16 experimental runs")
    add_image(doc, "ACTUAL IMAGE OF THE 16 DATASETS .png", IMAGES[8][1], 6.5)
    add_csv_table(doc, ROOT / "data" / "run_visual_labels.csv", "SEM visual labels")

    doc.add_heading("6. Methodology", 1)
    for s in [
        "Extract 16 runs; label SEM (Run 4 = only success)",
        "Phase 1: unknown position — GP + LOOCV on boundary circularity",
        "Phase 2: known (x,y) in 1500 um working area",
        "Compare with/without SEM images",
        "Scripts: phase1_model_actual.py, phase1_model.py",
    ]:
        doc.add_paragraph(s, style="List Number")

    doc.add_heading("7. Final results", 1)
    doc.add_paragraph("Phase 1 WITH images: 4 A, 150 us, 80 % (Run 4)")
    doc.add_paragraph("Phase 1 WITHOUT images: 6 A, 50 us, 64 % — NOT recommended")
    doc.add_paragraph("Phase 2 WITH images: 4 A at center; 3.5 A near struts")
    add_csv_table(doc, ROOT / "data" / "recommended_trials.csv", "Recommended trials")

    doc.add_heading("8. Conclusion", 1)
    doc.add_paragraph(
        "Use Run 4 parameters with fine feed and stable gap. Synthetic 1100 points are AI-generated, "
        "not independent lab data. Future: web tool for circularity heatmap over 1500x1500 um grid."
    )

    doc.save(OUTPUT)
    print(f"Saved: {OUTPUT}")


if __name__ == "__main__":
    build()
